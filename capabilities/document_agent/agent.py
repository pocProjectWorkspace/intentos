"""
IntentOS document_agent — Document creation, reading, and conversion

Primitive actions:
  - create_document: create a new .docx file with formatted content
  - append_content: add content to an existing .docx file
  - read_document: read a .docx or .pdf and return plain text
  - convert_document: convert between formats (docx/pdf → txt)
  - save_document: save/copy a document to a specified location

Default save location: ~/.intentos/workspace/outputs/
Uses python-docx for .docx, pypdf for PDF reading.

Follows SPEC.md: run() entry point, standard output dicts, dry_run support,
plain-language errors, audit metadata.
"""

from __future__ import annotations

import os
import shutil
import time
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _meta(
    files_affected: int = 0,
    bytes_affected: int = 0,
    duration_ms: int = 0,
    paths_accessed: list[str] | None = None,
) -> dict:
    return {
        "files_affected": files_affected,
        "bytes_affected": bytes_affected,
        "duration_ms": duration_ms,
        "paths_accessed": paths_accessed or [],
    }


def _error(code: str, message: str) -> dict:
    return {
        "status": "error",
        "error": {"code": code, "message": message},
        "metadata": _meta(),
    }


def _is_path_allowed(path: str, granted_paths: list[str]) -> bool:
    """Check if a path is within the granted paths."""
    resolved = str(Path(path).expanduser().resolve())
    for gp in granted_paths:
        gp_resolved = str(Path(gp).expanduser().resolve())
        if resolved == gp_resolved or resolved.startswith(gp_resolved + os.sep):
            return True
    return False


def _default_output_dir(context: dict) -> str:
    """Get the default output directory from context workspace."""
    workspace = context.get("workspace", "")
    if workspace:
        output_dir = os.path.join(workspace, "outputs")
    else:
        output_dir = os.path.join(
            os.path.expanduser("~"), ".intentos", "workspace", "outputs"
        )
    os.makedirs(output_dir, exist_ok=True)
    return output_dir


def _ensure_docx_extension(filename: str) -> str:
    """Ensure filename ends with .docx (handles .doc, no extension, etc.)."""
    base, ext = os.path.splitext(filename)
    if ext.lower() in (".doc", ".docx"):
        return base + ".docx"
    if not ext:
        return filename + ".docx"
    return filename


def _normalize_content(raw_content) -> str:
    """Normalize content from various input types to a plain string."""
    if isinstance(raw_content, dict):
        content = raw_content.get("extracted_data", "") or raw_content.get("content", "")
        if not content:
            parts = []
            for k, v in raw_content.items():
                if isinstance(v, str) and v:
                    parts.append(v)
            content = "\n\n".join(parts)
        return content
    elif isinstance(raw_content, list):
        parts = []
        for item in raw_content:
            if isinstance(item, dict):
                parts.append(" - ".join(str(v) for v in item.values() if v))
            else:
                parts.append(str(item))
        return "\n".join(parts)
    return str(raw_content).strip()


# ---------------------------------------------------------------------------
# create_document — create a new .docx with formatted content
# ---------------------------------------------------------------------------

def _create_document(params: dict, context: dict) -> dict:
    """Create a new .docx file with content."""
    t0 = time.monotonic()
    filename = params.get("filename", "").strip()
    raw_content = params.get("content", "")
    title = params.get("title", "").strip()
    save_path = params.get("save_path", "").strip()

    content = _normalize_content(raw_content)

    if not filename:
        return _error("MISSING_FILENAME", "I need a filename for the document")

    filename = _ensure_docx_extension(filename)

    # Determine output path
    if save_path:
        save_path = str(Path(save_path).expanduser().resolve())
        if os.path.isdir(save_path):
            output_path = os.path.join(save_path, filename)
        else:
            output_path = save_path
    else:
        output_dir = _default_output_dir(context)
        output_path = os.path.join(output_dir, filename)

    if context.get("dry_run"):
        return {
            "status": "success",
            "action_performed": f"Would create document '{filename}'",
            "result": {"preview": f"Create {filename} at {output_path}"},
            "metadata": _meta(paths_accessed=[output_path]),
        }

    try:
        doc = Document()

        # Add title if provided
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add date line
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        date_run.font.size = Pt(10)
        date_run.italic = True

        # Add content — split by double newlines into paragraphs
        if content:
            paragraphs = content.split("\n\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue
                # Handle single newlines within a paragraph as line breaks
                lines = para_text.split("\n")
                p = doc.add_paragraph()
                for j, line in enumerate(lines):
                    run = p.add_run(line.strip())
                    run.font.size = Pt(11)
                    if j < len(lines) - 1:
                        p.add_run("\n")

        # Add timestamp footer
        doc.add_paragraph()
        footer_text = f"Generated by IntentOS on {datetime.now().strftime('%Y-%m-%d at %H:%M')}"
        footer = doc.add_paragraph()
        run = footer.add_run(footer_text)
        run.font.size = Pt(8)
        run.italic = True

        # Save
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
    except PermissionError:
        return _error("PERMISSION_DENIED", "I don't have permission to save the document there")
    except Exception:
        return _error("CREATE_FAILED", "Something went wrong while creating the document")

    file_size = os.path.getsize(output_path) if os.path.exists(output_path) else 0
    elapsed = int((time.monotonic() - t0) * 1000)
    return {
        "status": "success",
        "action_performed": f"Created document '{filename}'",
        "result": {
            "path": output_path,
            "filename": filename,
            "size_bytes": file_size,
        },
        "metadata": _meta(
            files_affected=1,
            bytes_affected=file_size,
            duration_ms=elapsed,
            paths_accessed=[output_path],
        ),
    }


# ---------------------------------------------------------------------------
# append_content — add content to an existing .docx
# ---------------------------------------------------------------------------

def _append_content(params: dict, context: dict) -> dict:
    """Add content to an existing .docx file."""
    t0 = time.monotonic()
    filepath = params.get("path", "").strip()
    raw_content = params.get("content", "")
    heading = params.get("heading", "").strip()

    if not filepath:
        return _error("MISSING_PATH", "I need the path to the document")

    filepath = str(Path(filepath).expanduser().resolve())
    content = _normalize_content(raw_content)

    if not content:
        return _error("MISSING_CONTENT", "I need content to add to the document")

    if not os.path.exists(filepath):
        return _error("FILE_NOT_FOUND", "I couldn't find that document — it may have been moved or deleted")

    granted = context.get("granted_paths", [])
    if granted and not _is_path_allowed(filepath, granted):
        return _error("PATH_NOT_GRANTED", "I don't have access to that location")

    if context.get("dry_run"):
        return {
            "status": "success",
            "action_performed": f"Would add content to '{Path(filepath).name}'",
            "result": {"preview": f"Add content to {filepath}"},
            "metadata": _meta(paths_accessed=[filepath]),
        }

    try:
        doc = Document(filepath)

        if heading:
            doc.add_heading(heading, level=2)

        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.size = Pt(11)

        doc.save(filepath)
    except PermissionError:
        return _error("PERMISSION_DENIED", "I don't have permission to modify that document")
    except Exception:
        return _error("WRITE_FAILED", "Something went wrong while writing to the document")

    file_size = os.path.getsize(filepath) if os.path.exists(filepath) else 0
    elapsed = int((time.monotonic() - t0) * 1000)
    return {
        "status": "success",
        "action_performed": f"Added content to '{Path(filepath).name}'",
        "result": {
            "path": filepath,
            "filename": Path(filepath).name,
            "size_bytes": file_size,
        },
        "metadata": _meta(
            files_affected=1,
            bytes_affected=file_size,
            duration_ms=elapsed,
            paths_accessed=[filepath],
        ),
    }


# ---------------------------------------------------------------------------
# read_document — read .docx or .pdf and return plain text
# ---------------------------------------------------------------------------

def _read_document(params: dict, context: dict) -> dict:
    """Read a .docx or .pdf file and return its plain text content."""
    t0 = time.monotonic()
    filepath = params.get("path", "").strip()

    if not filepath:
        return _error("MISSING_PATH", "I need the path to the document")

    filepath = str(Path(filepath).expanduser().resolve())

    if not os.path.exists(filepath):
        return _error("FILE_NOT_FOUND", "I couldn't find that document — it may have been moved or deleted")

    granted = context.get("granted_paths", [])
    if granted and not _is_path_allowed(filepath, granted):
        return _error("PATH_NOT_GRANTED", "I don't have access to that location")

    ext = Path(filepath).suffix.lower()
    if ext not in (".docx", ".doc", ".pdf"):
        return _error("UNSUPPORTED_FORMAT", f"I can read .docx and .pdf files, but not '{ext}' files")

    if context.get("dry_run"):
        return {
            "status": "success",
            "action_performed": f"Would read '{Path(filepath).name}'",
            "result": {"preview": f"Read content from {filepath}"},
            "metadata": _meta(paths_accessed=[filepath]),
        }

    text = ""
    try:
        if ext in (".docx", ".doc"):
            doc = Document(filepath)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            text = "\n\n".join(paragraphs)
        elif ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(filepath)
            pages = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text.strip())
            text = "\n\n".join(pages)
    except PermissionError:
        return _error("PERMISSION_DENIED", "I don't have permission to read that document")
    except Exception:
        return _error("READ_FAILED", "Something went wrong while reading the document")

    file_size = os.path.getsize(filepath)
    elapsed = int((time.monotonic() - t0) * 1000)
    return {
        "status": "success",
        "action_performed": f"Read '{Path(filepath).name}' ({len(text):,} chars)",
        "result": {
            "path": filepath,
            "filename": Path(filepath).name,
            "content": text,
            "char_count": len(text),
            "format": ext.lstrip("."),
        },
        "metadata": _meta(
            files_affected=1,
            bytes_affected=file_size,
            duration_ms=elapsed,
            paths_accessed=[filepath],
        ),
    }


# ---------------------------------------------------------------------------
# convert_document — convert between formats (docx/pdf → txt)
# ---------------------------------------------------------------------------

def _convert_document(params: dict, context: dict) -> dict:
    """Convert a document between formats. Supports docx→txt and pdf→txt."""
    t0 = time.monotonic()
    filepath = params.get("path", "").strip()
    output_format = params.get("format", "txt").strip().lower().lstrip(".")

    if not filepath:
        return _error("MISSING_PATH", "I need the path to the document to convert")

    filepath = str(Path(filepath).expanduser().resolve())

    if not os.path.exists(filepath):
        return _error("FILE_NOT_FOUND", "I couldn't find that document")

    granted = context.get("granted_paths", [])
    if granted and not _is_path_allowed(filepath, granted):
        return _error("PATH_NOT_GRANTED", "I don't have access to that location")

    source_ext = Path(filepath).suffix.lower()
    supported_conversions = {
        (".docx", "txt"), (".doc", "txt"),
        (".pdf", "txt"),
    }

    if (source_ext, output_format) not in supported_conversions:
        return _error(
            "UNSUPPORTED_CONVERSION",
            f"I can't convert {source_ext} to .{output_format} — "
            f"supported: .docx→.txt, .pdf→.txt",
        )

    # Build output path in workspace
    base_name = Path(filepath).stem
    output_filename = f"{base_name}.{output_format}"
    output_dir = _default_output_dir(context)
    output_path = os.path.join(output_dir, output_filename)

    if context.get("dry_run"):
        return {
            "status": "success",
            "action_performed": f"Would convert '{Path(filepath).name}' to .{output_format}",
            "result": {"preview": f"Convert to {output_path}"},
            "metadata": _meta(paths_accessed=[filepath]),
        }

    # Extract text from source
    text = ""
    try:
        if source_ext in (".docx", ".doc"):
            doc = Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
        elif source_ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(filepath)
            pages = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text.strip())
            text = "\n\n".join(pages)
    except PermissionError:
        return _error("PERMISSION_DENIED", "I don't have permission to read that document")
    except Exception:
        return _error("CONVERT_FAILED", "Something went wrong while reading the source document")

    # Write output
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
    except PermissionError:
        return _error("PERMISSION_DENIED", "I don't have permission to save the converted file")
    except Exception:
        return _error("CONVERT_FAILED", "Something went wrong while saving the converted document")

    file_size = os.path.getsize(output_path) if os.path.exists(output_path) else 0
    elapsed = int((time.monotonic() - t0) * 1000)
    return {
        "status": "success",
        "action_performed": f"Converted '{Path(filepath).name}' to '{output_filename}'",
        "result": {
            "path": output_path,
            "filename": output_filename,
            "source": filepath,
            "size_bytes": file_size,
        },
        "metadata": _meta(
            files_affected=1,
            bytes_affected=file_size,
            duration_ms=elapsed,
            paths_accessed=[filepath, output_path],
        ),
    }


# ---------------------------------------------------------------------------
# save_document — save/copy a document to a specified location
# ---------------------------------------------------------------------------

def _save_document(params: dict, context: dict) -> dict:
    """Save or copy a document to a location."""
    t0 = time.monotonic()
    source = params.get("path", "").strip() or params.get("source", "").strip()
    destination = params.get("destination", "").strip()
    filename = params.get("filename", "").strip()

    if not source:
        return _error("MISSING_PATH", "I need the path to the source document")

    source = str(Path(source).expanduser().resolve())

    if not os.path.exists(source):
        return _error("FILE_NOT_FOUND", "I couldn't find that document")

    # Determine destination
    if destination:
        destination = str(Path(destination).expanduser().resolve())
        if os.path.isdir(destination):
            dest_filename = filename or Path(source).name
            destination = os.path.join(destination, dest_filename)
    elif filename:
        output_dir = _default_output_dir(context)
        destination = os.path.join(output_dir, _ensure_docx_extension(filename))
    else:
        destination = source

    if context.get("dry_run"):
        return {
            "status": "success",
            "action_performed": f"Would save document to '{Path(destination).name}'",
            "result": {"preview": f"Save to {destination}"},
            "metadata": _meta(paths_accessed=[source]),
        }

    try:
        if source != destination:
            os.makedirs(os.path.dirname(destination), exist_ok=True)
            shutil.copy2(source, destination)
    except PermissionError:
        return _error("PERMISSION_DENIED", "I don't have permission to save the document there")
    except Exception:
        return _error("SAVE_FAILED", "Something went wrong while saving the document")

    file_size = os.path.getsize(destination) if os.path.exists(destination) else 0
    elapsed = int((time.monotonic() - t0) * 1000)
    return {
        "status": "success",
        "action_performed": f"Saved document to '{Path(destination).name}'",
        "result": {
            "path": destination,
            "filename": Path(destination).name,
            "size_bytes": file_size,
        },
        "metadata": _meta(
            files_affected=1,
            bytes_affected=file_size,
            duration_ms=elapsed,
            paths_accessed=[source, destination],
        ),
    }


# ---------------------------------------------------------------------------
# Action registry and entry point
# ---------------------------------------------------------------------------

_ACTIONS = {
    "create_document": _create_document,
    "append_content": _append_content,
    "write_content": _append_content,  # backward compat alias
    "read_document": _read_document,
    "convert_document": _convert_document,
    "save_document": _save_document,
}


def run(input: dict) -> dict:
    action = input.get("action")
    params = input.get("params", {})
    context = input.get("context", {})

    handler = _ACTIONS.get(action)
    if handler is None:
        return _error("UNKNOWN_ACTION", f"I don't know how to do '{action}'")

    try:
        return handler(params, context)
    except Exception:
        return _error("AGENT_CRASH", "Something went wrong running that operation")
